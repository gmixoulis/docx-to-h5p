#!/usr/bin/env python3
"""
Batch DOCX to H5P Converter
Processes all DOCX files in a folder and creates separate output directories
Supports: Multiple Choice, True/False, and Crossword Puzzles
"""

import json
import re
import uuid
from docx import Document
import os
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path


class CompleteH5PConverter:
    def __init__(self):
        self.multiple_choice_questions = []
        self.true_false_questions = []
        self.crossword_puzzles = []
        self.images = {}
        self.shared_image_id = None


    def reset(self):
        """Reset all lists for processing a new document."""
        self.multiple_choice_questions = []
        self.true_false_questions = []
        self.crossword_puzzles = []
        self.images = {}
        self.shared_image_id = None


    def extract_images_from_docx(self, docx_path: str) -> Dict[str, Dict]:
        """Extract all images from DOCX with their metadata."""
        doc = Document(docx_path)
        images_info = {}

        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    image_bytes = image_part.blob

                    content_type = image_part.content_type
                    ext = content_type.split('/')[-1]
                    if ext == 'jpeg':
                        ext = 'jpg'

                    image_name = f"image_{rel_id}.{ext}"

                    images_info[rel_id] = {
                        "name": image_name,
                        "bytes": image_bytes,
                        "mime": content_type,
                        "size": len(image_bytes),
                        "width": 600,
                        "height": 400
                    }
        except Exception as e:
            print(f"  Ã¢Å¡ Ã¯Â¸Â  Warning: Could not extract images: {e}")

        return images_info


    def find_image_near_paragraph(self, doc: Document, para_index: int) -> Optional[str]:
        """Find if there's an image near a specific paragraph."""
        search_range = range(max(0, para_index - 5), min(len(doc.paragraphs), para_index + 3))

        for idx in search_range:
            try:
                para = doc.paragraphs[idx]
                if para._element.xpath('.//pic:pic'):
                    for run in para.runs:
                        blips = run._element.xpath('.//a:blip')
                        if blips:
                            embed = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed:
                                return embed
            except:
                continue

        return None


    # ========================================================================
    # TRUE/FALSE QUESTION EXTRACTION
    # ========================================================================

    def extract_text_with_formatting(self, paragraph):
        """
        Extract text from paragraph while tracking which parts are bold
        Returns: list of tuples (text, is_bold)
        """
        result = []
        for run in paragraph.runs:
            if run.text.strip():
                result.append((run.text, run.bold))
        return result

    def is_true_false_question(self, text):
        """
        Detect if the text represents a True/False question
        """
        clean_text = text.strip()

        # Check if text contains True or False at the end
        patterns = [
            r'\b(true|false)\s*$',
            r'\*\*(true|false)\*\*\s*$',
        ]

        for pattern in patterns:
            if re.search(pattern, clean_text, re.IGNORECASE):
                return True

        return False

    def parse_true_false_from_runs(self, paragraph):
        """
        Parse True/False question from paragraph runs
        Identifies which answer (True/False) is bold
        Returns: (question_text, correct_answer)
        """
        formatted_parts = self.extract_text_with_formatting(paragraph)

        full_text = ''.join([text for text, _ in formatted_parts])

        # Find True/False at the end
        true_false_pattern = r'\b(True|False)\b'
        matches = list(re.finditer(true_false_pattern, full_text, re.IGNORECASE))

        if not matches:
            return None, None

        # Get the last True/False match
        last_match = matches[-1]
        answer_text = last_match.group(1)
        answer_start = last_match.start()

        # Check if this answer is bold
        current_pos = 0
        is_bold = False

        for text, bold in formatted_parts:
            if current_pos <= answer_start < current_pos + len(text):
                is_bold = bold
                break
            current_pos += len(text)

        # Extract question without the answer
        question_text = full_text[:answer_start].strip()

        # If bold, this is the correct answer; otherwise it's the opposite
        correct_answer = answer_text.lower() if is_bold else ('false' if answer_text.lower() == 'true' else 'true')

        return question_text, correct_answer

    def extract_true_false_questions(self, doc: Document) -> None:
        """Extract True/False questions from document."""

        in_true_false_section = False

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            if not text:
                continue

            # Detect True/False section header
            if 'true or false' in text.lower() or 'true/false' in text.lower():
                in_true_false_section = True
                print(f"\n  Ã°Å¸â€œÂ Found True/False section: {text}")
                continue

            # Exit True/False section if we hit a new activity or section
            if in_true_false_section:
                if (text.startswith('**') and ('activity' in text.lower() or 'quiz' in text.lower()) 
                    and 'true' not in text.lower() and 'false' not in text.lower()):
                    in_true_false_section = False
                    print(f"  Ã°Å¸â€œÂ End of True/False section\n")
                    continue

            # Process True/False questions
            if in_true_false_section and self.is_true_false_question(text):
                question_text, correct_answer = self.parse_true_false_from_runs(para)

                if question_text and correct_answer:
                    self.true_false_questions.append({
                        'question': question_text,
                        'correct_answer': correct_answer
                    })
                    print(f"    Ã¢Å“â€œ T/F Question: {question_text[:60]}... [{correct_answer.upper()}]")


    # ========================================================================
    # MULTIPLE CHOICE QUESTION EXTRACTION
    # ========================================================================

    def extract_multiple_choice_questions(self, doc: Document, docx_path: str) -> None:
        """Extract multiple choice questions from document."""

        self.images = self.extract_images_from_docx(docx_path)

        in_mc_section = False
        section_image_id = None
        questions_in_current_section = []

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            if not text:
                continue

            # Detect Activity 1 section (multiple choice) - usually quiz section
            if re.search(r'Activity\s*1.*quiz', text, re.IGNORECASE) or re.search(r'quiz.*question', text, re.IGNORECASE):
                in_mc_section = True
                section_image_id = self.find_image_near_paragraph(doc, idx)
                if section_image_id:
                    print(f"  Ã°Å¸â€“Â¼Ã¯Â¸Â  Found shared image for quiz section")
                continue

            # Exit MC section when hitting Activity 2 or other sections
            if re.search(r'Activity\s*[2-9]', text, re.IGNORECASE) and in_mc_section:
                in_mc_section = False
                if section_image_id and questions_in_current_section:
                    self.shared_image_id = section_image_id
                continue

            if not in_mc_section:
                continue

            lines = text.split('\n')
            first_line = lines[0] if lines else ""
            question_match = re.match(r'^(\d+)\.\s*(.+?)\?\s*$', first_line)

            if question_match and len(lines) > 1:
                question_num = question_match.group(1)
                question_text = question_match.group(2).strip() + "?"

                options = []
                option_runs_map = {}

                current_line = 0
                current_pos = 0

                for run in para.runs:
                    run_text = run.text
                    while current_line < len(lines):
                        line_text = lines[current_line]
                        if current_pos < len(line_text):
                            break
                        current_pos = 0
                        current_line += 1

                    if current_line < len(lines):
                        option_match = re.match(r'^([A-D])\.\s', lines[current_line])
                        if option_match:
                            option_letter = option_match.group(1)
                            if option_letter not in option_runs_map:
                                option_runs_map[option_letter] = []
                            option_runs_map[option_letter].append(run)

                    current_pos += len(run_text)

                for line in lines[1:]:
                    option_match = re.match(r'^([A-D])\.\s+(.+)$', line)
                    if option_match:
                        option_letter = option_match.group(1)
                        option_text = option_match.group(2).strip().rstrip('.')

                        is_correct = False
                        if option_letter in option_runs_map:
                            is_correct = any(run.bold for run in option_runs_map[option_letter] if run.text.strip())

                        options.append({
                            "text": option_text,
                            "correct": is_correct
                        })

                if options:
                    question_image_id = self.find_image_near_paragraph(doc, idx)
                    if not question_image_id:
                        question_image_id = section_image_id

                    self.multiple_choice_questions.append({
                        "question": question_text,
                        "options": options,
                        "image_id": question_image_id
                    })

                    questions_in_current_section.append(len(self.multiple_choice_questions) - 1)


    # ========================================================================
    # CROSSWORD PUZZLE EXTRACTION
    # ========================================================================

    def extract_crossword_puzzles(self, doc: Document) -> None:
        """Extract crossword puzzles - ULTRA ROBUST VERSION."""

        crosswords = []
        current_crossword = None
        current_orientation = None

        in_crossword_section = False
        collecting_clues = False
        clue_counter = 1

        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()

            if not text:
                i += 1
                continue

            # Detect crossword section start
            crossword_match = re.search(r'Activity\s*3,?\s*Part\s+(I+)[-\s]*Crossword\s*Puzzle:?(.*)$', text, re.IGNORECASE)

            if crossword_match:
                if current_crossword and current_crossword.get('clues'):
                    crosswords.append(current_crossword)

                part_num = crossword_match.group(1)
                title_suffix = crossword_match.group(2).strip()

                full_title = f"Activity 3, Part {part_num} - Crossword Puzzle"
                if title_suffix:
                    full_title += f": {title_suffix}"

                current_crossword = {
                    "title": full_title,
                    "clues": []
                }

                in_crossword_section = True
                collecting_clues = False
                current_orientation = None
                clue_counter = 1

                print(f"\n  Ã°Å¸Â§Â© Found: {full_title}")
                i += 1
                continue

            # Exit if we hit Unit 2 or Activity 4
            if re.search(r'Unit\s*2', text, re.IGNORECASE) or re.search(r'Activity\s*4', text, re.IGNORECASE):
                if current_crossword and current_crossword.get('clues'):
                    crosswords.append(current_crossword)
                    current_crossword = None
                in_crossword_section = False
                i += 1
                continue

            if not in_crossword_section or not current_crossword:
                i += 1
                continue

            # Look for "Clues:" marker
            if re.match(r'^Clues?:?\s*$', text, re.IGNORECASE):
                collecting_clues = True
                i += 1
                continue

            # Detect orientation
            if re.search(r'^(Across|Down)', text, re.IGNORECASE):
                orientation_match = re.match(r'^(Across|Down)', text, re.IGNORECASE)
                current_orientation = orientation_match.group(1).lower()
                collecting_clues = True
                clue_counter = 1
                print(f"    Ã°Å¸â€œÂ {current_orientation.title()} clues:")

                rest_of_text = text[len(orientation_match.group(0)):].strip()

                if rest_of_text:
                    lines = rest_of_text.split('\n')

                    for line in lines:
                        clue_pattern = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
                        if clue_pattern:
                            clue_num = clue_pattern.group(1)
                            clue_text = clue_pattern.group(2).strip()

                            bold_answers = [run.text.strip().upper() for run in para.runs 
                                          if run.bold and run.text.strip() and len(run.text.strip()) > 2]

                            if bold_answers:
                                for answer in bold_answers:
                                    answer_clean = ' '.join(answer.split())

                                    current_crossword["clues"].append({
                                        "orientation": current_orientation,
                                        "clue": clue_text,
                                        "answer": answer_clean
                                    })
                                    print(f"      Ã¢Å“â€œ Clue {clue_num}: {clue_text[:40]}... Ã¢â€ â€™ {answer_clean}")

                i += 1
                continue

            if not collecting_clues or not current_orientation:
                i += 1
                continue

            # Pattern 1: Numbered clues
            clue_pattern = re.match(r'^(\d+)\.\s+(.+?)(?:\s*\(([^)]+)\))?\s*$', text)

            if clue_pattern:
                clue_num = clue_pattern.group(1)
                clue_text = clue_pattern.group(2).strip()

                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    next_text = next_para.text.strip()

                    if re.match(r'^[A-Z][A-Z\s-]+$', next_text):
                        answer = ' '.join(next_text.split())

                        current_crossword["clues"].append({
                            "orientation": current_orientation,
                            "clue": clue_text,
                            "answer": answer
                        })
                        print(f"      Ã¢Å“â€œ Clue {clue_num}: {clue_text[:40]}... Ã¢â€ â€™ {answer}")
                        i += 2
                        continue

            # Pattern 2: Clues without numbers
            else:
                if '(' in text and ')' in text:
                    clue_text = text.strip()

                    if i + 1 < len(doc.paragraphs):
                        next_para = doc.paragraphs[i + 1]
                        next_text = next_para.text.strip()

                        bold_parts = [run.text.strip().upper() for run in next_para.runs 
                                    if run.bold and run.text.strip()]

                        if bold_parts and re.match(r'^[A-Z][A-Z\s-]+$', bold_parts[0]):
                            answer = ' '.join(bold_parts[0].split())

                            current_crossword["clues"].append({
                                "orientation": current_orientation,
                                "clue": clue_text,
                                "answer": answer
                            })
                            print(f"      Ã¢Å“â€œ Clue {clue_counter}: {clue_text[:40]}... Ã¢â€ â€™ {answer}")
                            clue_counter += 1
                            i += 2
                            continue

            i += 1

        # Save final crossword
        if current_crossword and current_crossword.get('clues'):
            crosswords.append(current_crossword)

        self.crossword_puzzles = crosswords


    # ========================================================================
    # H5P JSON GENERATION
    # ========================================================================

    def create_h5p_question_json(self, question_data: Dict) -> Dict:
        """Create H5P multiple choice JSON."""

        answers = []
        for option in question_data["options"]:
            answer = {
                "correct": option["correct"],
                "text": f"<div>{option['text']}</div>\n",
                "tipsAndFeedback": {
                    "tip": "",
                    "chosenFeedback": "",
                    "notChosenFeedback": ""
                }
            }
            answers.append(answer)

        question_json = {
            "answers": answers,
            "UI": {
                "showSolutionButton": "Show solution",
                "tryAgainButton": "Retry",
                "checkAnswerButton": "Check",
                "tipsLabel": "Show tip",
                "scoreBarLabel": "You got :num out of :total points",
                "tipAvailable": "Tip available",
                "feedbackAvailable": "Feedback available",
                "readFeedback": "Read feedback",
                "wrongAnswer": "Wrong answer",
                "correctAnswer": "Correct answer",
                "shouldCheck": "Should have been checked",
                "shouldNotCheck": "Should not have been checked",
                "noInput": "Please answer before viewing the solution",
                "submitAnswerButton": "Submit",
                "a11yCheck": "Check the answers.",
                "a11yShowSolution": "Show the solution.",
                "a11yRetry": "Retry the task."
            },
            "question": f"<p>{question_data['question']}</p>\n",
            "behaviour": {
                "enableRetry": True,
                "enableSolutionsButton": True,
                "singlePoint": True,
                "randomAnswers": True,
                "showSolutionsRequiresInput": True,
                "type": "auto",
                "confirmCheckDialog": False,
                "confirmRetryDialog": False,
                "autoCheck": False,
                "passPercentage": 100,
                "showScorePoints": True,
                "enableCheckButton": True
            },
            "confirmCheck": {
                "header": "Finish ?",
                "body": "Are you sure you wish to finish ?",
                "cancelLabel": "Cancel",
                "confirmLabel": "Finish"
            },
            "confirmRetry": {
                "header": "Retry ?",
                "body": "Are you sure you wish to retry ?",
                "cancelLabel": "Cancel",
                "confirmLabel": "Confirm"
            },
            "overallFeedback": [
                {"from": 0, "to": 0, "feedback": "Wrong"},
                {"from": 1, "to": 99, "feedback": "Almost!"},
                {"from": 100, "to": 100, "feedback": "Correct!"}
            ]
        }

        image_id = question_data.get("image_id")
        if image_id and image_id in self.images:
            img_info = self.images[image_id]

            question_json["media"] = {
                "type": {
                    "params": {
                        "contentName": "Image",
                        "file": {
                            "path": f"images/{img_info['name']}",
                            "mime": img_info["mime"],
                            "width": img_info["width"],
                            "height": img_info["height"],
                            "copyright": {"license": "U"}
                        },
                        "alt": f"Image for question",
                        "decorative": False
                    },
                    "library": "H5P.Image 1.1",
                    "subContentId": str(uuid.uuid4()),
                    "metadata": {
                        "title": "Question Image",
                        "license": "U",
                        "contentType": "Image"
                    }
                },
                "disableImageZooming": True
            }
        else:
            question_json["media"] = {"disableImageZooming": True}

        return question_json

    def create_h5p_true_false_json(self, question_data: Dict) -> Dict:
        """Create H5P True/False JSON."""

        content = {
            "media": {
                "type": {},
                "disableImageZooming": False
            },
            "correct": question_data['correct_answer'],
            "l10n": {
                "trueText": "True",
                "falseText": "False",
                "score": "You got @score of @total points",
                "checkAnswer": "Check",
                "showSolutionButton": "Show solution",
                "tryAgain": "Retry",
                "wrongAnswerMessage": "Wrong answer",
                "correctAnswerMessage": "Correct answer",
                "scoreBarLabel": "You got :num out of :total points",
                "submitAnswer": "Submit",
                "a11yCheck": "Check the answers. The responses will be marked as correct, incorrect, or unanswered.",
                "a11yShowSolution": "Show the solution. The task will be marked with its correct solution.",
                "a11yRetry": "Retry the task. Reset all responses and start the task over again."
            },
            "behaviour": {
                "enableRetry": True,
                "enableSolutionsButton": True,
                "confirmCheckDialog": False,
                "confirmRetryDialog": False,
                "autoCheck": False,
                "enableCheckButton": True
            },
            "confirmCheck": {
                "header": "Finish ?",
                "body": "Are you sure you wish to finish ?",
                "cancelLabel": "Cancel",
                "confirmLabel": "Finish"
            },
            "confirmRetry": {
                "header": "Retry ?",
                "body": "Are you sure you wish to retry ?",
                "cancelLabel": "Cancel",
                "confirmLabel": "Confirm"
            },
            "question": f"<p>{question_data['question']}</p>\n"
        }

        return content

    def create_h5p_crossword_json(self, crossword_data: Dict) -> Dict:
        """Create H5P crossword JSON."""
        words = []

        for clue_data in crossword_data["clues"]:
            word_entry = {
                "fixWord": False,
                "orientation": clue_data["orientation"],
                "clue": clue_data["clue"],
                "answer": clue_data["answer"]
            }
            words.append(word_entry)

        return {
            "words": words,
            "overallFeedback": [{"from": 0, "to": 100}],
            "theme": {
                "backgroundColor": "#222b46",
                "gridColor": "#031928",
                "cellBackgroundColor": "#ffffff",
                "cellColor": "#000000",
                "clueIdColor": "#606060",
                "cellBackgroundColorHighlight": "#5c9ba9",
                "cellColorHighlight": "#031928",
                "clueIdColorHighlight": "#e0e0e0"
            },
            "behaviour": {
                "enableInstantFeedback": False,
                "scoreWords": True,
                "applyPenalties": False,
                "enableRetry": True,
                "enableSolutionsButton": True
            },
            "l10n": {
                "across": "Across",
                "down": "Down",
                "checkAnswer": "Check",
                "tryAgain": "Retry",
                "showSolution": "Show solution",
                "couldNotGenerateCrossword": "Could not generate crossword.",
                "couldNotGenerateCrosswordTooFewWords": "Need at least two words.",
                "probematicWords": "Problematic word(s): @words",
                "extraClue": "Extra clue",
                "closeWindow": "Close window",
                "submitAnswer": "Submit"
            },
            "a11y": {
                "crosswordGrid": "Crossword grid.",
                "column": "Column",
                "row": "Row",
                "across": "Across",
                "down": "Down",
                "empty": "Empty",
                "resultFor": "Result for: @clue",
                "correct": "Correct",
                "wrong": "Wrong",
                "point": "point",
                "solutionFor": "Solution: @solution",
                "check": "Check",
                "showSolution": "Show solution",
                "retry": "Retry",
                "yourResult": "You got @score out of @total points"
            },
            "taskDescription": f"<p>{crossword_data['title']}</p>\n"
        }


    # ========================================================================
    # FILE SAVING
    # ========================================================================

    def save_images(self, output_dir: str) -> None:
        """Save extracted images."""
        if not self.images:
            return

        images_dir = os.path.join(output_dir, "images")
        os.makedirs(images_dir, exist_ok=True)

        for img_id, img_info in self.images.items():
            img_path = os.path.join(images_dir, img_info["name"])
            with open(img_path, 'wb') as f:
                f.write(img_info["bytes"])
            print(f"  Ã¢Å“â€œ Saved: {img_info['name']}")


    # ========================================================================
    # MAIN PROCESSING
    # ========================================================================

    def process_docx(self, docx_path: str, output_dir: str = "output") -> None:
        """Main processing function."""
        print(f"\nÃ°Å¸â€œâ€ž Processing: {docx_path}")
        print("=" * 70)

        doc = Document(docx_path)

        print("\nÃ°Å¸â€œÂ Extracting multiple choice questions...")
        self.extract_multiple_choice_questions(doc, docx_path)
        print(f"  Ã¢Å“â€œ Found: {len(self.multiple_choice_questions)} questions")

        print("\nÃ¢Å“â€œ/Ã¢Å“â€” Extracting True/False questions...")
        self.extract_true_false_questions(doc)
        print(f"  Ã¢Å“â€œ Found: {len(self.true_false_questions)} questions")

        print("\nÃ°Å¸Â§Â© Extracting crossword puzzles...")
        self.extract_crossword_puzzles(doc)
        print(f"  Ã¢Å“â€œ Found: {len(self.crossword_puzzles)} crosswords")

        for i, cw in enumerate(self.crossword_puzzles, 1):
            across_count = sum(1 for c in cw['clues'] if c['orientation'] == 'across')
            down_count = sum(1 for c in cw['clues'] if c['orientation'] == 'down')
            print(f"    {i}. {cw['title'][:60]}")
            print(f"       Ã¢â€ â€™ {len(cw['clues'])} total ({across_count} across, {down_count} down)")

        os.makedirs(output_dir, exist_ok=True)

        if self.images:
            print(f"\nÃ°Å¸â€™Â¾ Saving images...")
            self.save_images(output_dir)

        print("\nÃ°Å¸â€œÂ Generating JSON files...")

        # Save Multiple Choice questions
        for i, question_data in enumerate(self.multiple_choice_questions, 1):
            question_json = self.create_h5p_question_json(question_data)

            filename = f"question_{i}.json"
            filepath = os.path.join(output_dir, filename)

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(question_json, f, indent=2, ensure_ascii=False)

            correct_answer = next((opt["text"] for opt in question_data["options"] if opt["correct"]), "N/A")
            print(f"  Ã¢Å“â€œ {filename}: {question_data['question'][:50]}...")

        # Save True/False questions
        for i, tf_data in enumerate(self.true_false_questions, 1):
            tf_json = self.create_h5p_true_false_json(tf_data)

            filename = f"true_false_{i}.json"
            filepath = os.path.join(output_dir, filename)

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(tf_json, f, indent=2, ensure_ascii=False)

            print(f"  Ã¢Å“â€œ {filename}: {tf_data['question'][:50]}... [{tf_data['correct_answer'].upper()}]")

        # Save Crosswords
        for i, crossword_data in enumerate(self.crossword_puzzles, 1):
            crossword_json = self.create_h5p_crossword_json(crossword_data)

            filename = f"crossword_{i}.json"
            filepath = os.path.join(output_dir, filename)

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(crossword_json, f, indent=2, ensure_ascii=False)

            across_count = sum(1 for c in crossword_data['clues'] if c['orientation'] == 'across')
            down_count = sum(1 for c in crossword_data['clues'] if c['orientation'] == 'down')
            print(f"  Ã¢Å“â€œ {filename}")
            print(f"    Ã¢â€ â€™ Across: {across_count}, Down: {down_count}")

        print(f"\n{'='*70}")
        print("Ã¢Å“â€¦ Complete! Summary:")
        print(f"  Multiple Choice: {len(self.multiple_choice_questions)}")
        print(f"  True/False: {len(self.true_false_questions)}")
        print(f"  Crosswords: {len(self.crossword_puzzles)}")
        print(f"  Output: {output_dir}/")
        print('='*70)


def process_folder(input_folder: str = "english_docs") -> None:
    """
    Process all DOCX files in a folder.
    Creates separate output folders for each DOCX file.

    Args:
        input_folder: Folder containing DOCX files (default: "english_docs")
    """
    # Check if input folder exists
    if not os.path.exists(input_folder):
        print(f"Ã¢ÂÅ’ Error: Folder '{input_folder}' not found!")
        print(f"\nPlease create the folder '{input_folder}' and add your DOCX files.")
        return

    # Find all DOCX files in the folder
    docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx') and not f.startswith('~')]

    if not docx_files:
        print(f"Ã¢ÂÅ’ Error: No DOCX files found in '{input_folder}' folder!")
        return

    print("\n" + "=" * 70)
    print("Ã°Å¸Å¡â‚¬ BATCH DOCX to H5P CONVERTER")
    print("=" * 70)
    print(f"\nÃ°Å¸â€œÂ Input folder: {input_folder}")
    print(f"Ã°Å¸â€œÅ  Found {len(docx_files)} DOCX file(s):\n")

    for i, file in enumerate(docx_files, 1):
        print(f"  {i}. {file}")

    print("\n" + "=" * 70)

    # Process each DOCX file
    converter = CompleteH5PConverter()

    for idx, docx_file in enumerate(docx_files, 1):
        # Reset converter for new document
        converter.reset()

        # Create paths
        docx_path = os.path.join(input_folder, docx_file)

        # Get filename without extension for output folder name
        output_folder_name = Path(docx_file).stem  # e.g., "Activities-Module-1"
        output_dir = output_folder_name  # Create folder with same name as DOCX

        print(f"\n\n{'='*70}")
        print(f"Ã°Å¸â€œâ€¹ Processing file {idx}/{len(docx_files)}: {docx_file}")
        print(f"Ã°Å¸â€œÂ Output folder: {output_dir}/")
        print('='*70)

        try:
            # Process the DOCX file
            converter.process_docx(docx_path, output_dir)

        except Exception as e:
            print(f"\nÃ¢ÂÅ’ Error processing {docx_file}: {e}")
            import traceback
            traceback.print_exc()
            continue

    # Final summary
    print("\n\n" + "=" * 70)
    print("Ã°Å¸Å½â€° BATCH PROCESSING COMPLETE!")
    print("=" * 70)
    print(f"\nÃ¢Å“â€¦ Successfully processed {len(docx_files)} file(s)")
    print(f"\nÃ°Å¸â€œÂ Output structure:")
    for docx_file in docx_files:
        folder_name = Path(docx_file).stem
        print(f"\n  {folder_name}/")
        print(f"    Ã¢â€Å“Ã¢â€â‚¬Ã¢â€â‚¬ question_*.json")
        print(f"    Ã¢â€Å“Ã¢â€â‚¬Ã¢â€â‚¬ true_false_*.json")
        print(f"    Ã¢â€Å“Ã¢â€â‚¬Ã¢â€â‚¬ crossword_*.json")
        print(f"    Ã¢â€â€Ã¢â€â‚¬Ã¢â€â‚¬ images/ (if any)")
    print("\n" + "=" * 70)


def main():
    import sys

    if len(sys.argv) > 1:
        # Single file mode (backward compatibility)
        docx_path = sys.argv[1]

        if not os.path.exists(docx_path):
            print(f"Ã¢ÂÅ’ Error: File not found: {docx_path}")
            return

        converter = CompleteH5PConverter()
        converter.process_docx(docx_path)
    else:
        # Batch folder mode (new default behavior)
        print("Complete DOCX to H5P Converter - Batch Mode")
        print("=" * 60)
        print("\nUsage:")
        print("  1. Batch mode: python batch_h5p_converter.py")
        print("     Processes all DOCX files in 'english_docs' folder")
        print("\n  2. Single file: python batch_h5p_converter.py <docx_file>")
        print("     Processes a single DOCX file")
        print("\nBatch mode features:")
        print("  Ã¢â‚¬Â¢ Processes all .docx files in 'english_docs' folder")
        print("  Ã¢â‚¬Â¢ Creates separate output folders for each DOCX")
        print("  Ã¢â‚¬Â¢ Output folder named after the DOCX file")
        print("\nExample folder structure:")
        print("  english_docs/")
        print("    Ã¢â€Å“Ã¢â€â‚¬Ã¢â€â‚¬ Activities-Module-1.docx")
        print("    Ã¢â€Å“Ã¢â€â‚¬Ã¢â€â‚¬ Activities-Module-2.docx")
        print("    Ã¢â€â€Ã¢â€â‚¬Ã¢â€â‚¬ Activities-Module-3.docx")
        print("\n  Output:")
        print("    Activities-Module-1/")
        print("      Ã¢â€Å“Ã¢â€â‚¬Ã¢â€â‚¬ question_1.json")
        print("      Ã¢â€Å“Ã¢â€â‚¬Ã¢â€â‚¬ true_false_1.json")
        print("      Ã¢â€â€Ã¢â€â‚¬Ã¢â€â‚¬ crossword_1.json")
        print("    Activities-Module-2/")
        print("      Ã¢â€â€Ã¢â€â‚¬Ã¢â€â‚¬ ...")
        print("    Activities-Module-3/")
        print("      Ã¢â€â€Ã¢â€â‚¬Ã¢â€â‚¬ ...")
        print("\n" + "=" * 60)

        input_folder = "english_docs"
        process_folder(input_folder)


if __name__ == "__main__":
    main()